from docx import Document

# Read content from the source text file (e.g., notepad.txt)
source_text_file = 'outputs.txt'
with open(source_text_file, 'r') as file:
        content = file.read()

        # Create a Word document
        doc = Document()
        doc.add_heading('Notepad Content', level=1)
        doc.add_paragraph(content)

        # Save the Word document
        output_docx_path = 'notepad_content.docx'
        doc.save(output_docx_path)
        print(f"Word document saved at: {output_docx_path}")
